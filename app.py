from flask import Flask, render_template, request, send_file
from docx import Document
import requests
import os
import re
import json
from dotenv import load_dotenv

# -------------------------------------------------------
# 1) Load environment variables (.env)
# -------------------------------------------------------
load_dotenv()

# -------------------------------------------------------
# 2) Groq client init → CLEAN & WORKING (proxy hack hata diya)
# -------------------------------------------------------
from groq import Groq

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
groq_client = None

if GROQ_API_KEY and GROQ_API_KEY.strip():
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
        print("Groq Loaded Successfully")
    except Exception as e:
        print("Groq Init Error:", str(e))
        groq_client = None
else:
    print("GROQ_API_KEY missing or empty")

# -------------------------------------------------------
# 3) Flask app + LanguageTool URL
# -------------------------------------------------------
app = Flask(__name__)
LT_API_URL = "https://api.languagetool.org/v2/check"

# -------------------------------------------------------
# 4) Load Black’s Law Dictionary JSON
# -------------------------------------------------------
try:
    with open("blacklaw_terms.json", "r", encoding="utf-8") as f:
        BLACKLAW = json.load(f)
except Exception:
    BLACKLAW = {}

def normalize_key(word: str) -> str:
    return re.sub(r"[^a-z\s]", "", word.lower().strip())

# -------------------------------------------------------
# 5) Legal fix rules
# -------------------------------------------------------
LEGAL_FIX = {
    "suo moto": "suo motu",
    "prima facia": "prima facie",
    "mens reaa": "mens rea",
    "ratio decedendi": "ratio decidendi",
}

IGNORE_WORDS = {
    "was", "the", "is", "and", "to", "in", "for",
    "of", "at", "by", "on", "with"
}

def is_reference_like(line: str) -> bool:
    s = line.strip()
    return (
        not s
        or "http" in s
        or "www." in s
        or "doi" in s.lower()
        or re.match(r"^\[\d+\]$", s)
        or re.match(r"^\(.+\d{4}.*\)$", s)
    )

# -------------------------------------------------------
# 6) LanguageTool API
# -------------------------------------------------------
def lt_check_sentence(sentence: str) -> dict:
    try:
        data = {"text": sentence, "language": "en-US"}
        r = requests.post(LT_API_URL, data=data, timeout=8)
        return r.json()
    except Exception:
        return {"matches": []}

# -------------------------------------------------------
# 7) Groq check (LLM grammar) + SAFE FIX
# -------------------------------------------------------
def groq_check(sentence: str, lt_wrong_words: list) -> list:
    if (not groq_client) or (not lt_wrong_words):
        return []

    lt_wrong_words = [
        w for w in lt_wrong_words if w.lower() not in IGNORE_WORDS
    ]
    if not lt_wrong_words:
        return []

    prompt = f"""
You are a hybrid LEGAL + GRAMMAR correction engine.
RULES:
1. Only correct words that appear in: {lt_wrong_words}
2. Never correct helper words: {list(IGNORE_WORDS)}
3. Apply Black's Law correction:
   - suo moto → suo motu
   - prima facia → prima facie
   - ratio decedendi → ratio decidendi
   - mens reaa → mens rea
4. Do NOT explain. Output ONLY a pure JSON array.

Sentence:
\"\"\"{sentence}\"\"\""""

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=512
        )
        raw = response.choices[0].message.content or ""
        match = re.search(r"\[.*?\]", raw, re.DOTALL)
        if not match:
            return []
        return json.loads(match.group(0))
    except Exception as e:
        print("GROQ ERROR:", e)
        return []

# -------------------------------------------------------
# 8) Legal phrase detection
# -------------------------------------------------------
def detect_legal(sentence: str):
    results = []
    for wrong, correct in LEGAL_FIX.items():
        if re.search(rf"\b{re.escape(wrong)}\b", sentence, re.IGNORECASE):
            meaning = BLACKLAW.get(normalize_key(correct), "")
            results.append((wrong, correct, meaning))
    return results

# -------------------------------------------------------
# 9) Build highlighted HTML (contains SAFE FIX)
# -------------------------------------------------------
def process_text_line_by_line(text: str) -> str:
    lines = text.split("\n")
    final_html = []
    for line in lines:
        if not line.strip():
            final_html.append("<p></p>")
            continue
        if is_reference_like(line):
            final_html.append(f"<p>{line}</p>")
            continue

        working = line
        html_line = line

        lt_res = lt_check_sentence(working)
        lt_wrong_words = []
        for m in lt_res.get("matches", []):
            wrong = working[m["offset"]:m["offset"] + m["length"]]
            if wrong.strip() and wrong.lower() not in IGNORE_WORDS:
                lt_wrong_words.append(wrong)

        legal_hits = detect_legal(working)
        groq_hits = groq_check(working, lt_wrong_words)

        combined = {}
        for wrong, correct, meaning in legal_hits:
            combined.setdefault(wrong.lower(), {"black": correct, "groq": None, "meaning": meaning})

        for g in groq_hits:
            if isinstance(g, dict):
                wrong = (g.get("wrong") or "").strip()
                suggestion = (g.get("suggestion") or "").strip()
                if wrong and suggestion:
                    key = wrong.lower()
                    if key not in combined:
                        combined[key] = {"black": None, "groq": None, "meaning": ""}
                    combined[key]["groq"] = suggestion

        for wrong_lower, data in combined.items():
            wrong_original = next(
                (w for w in re.findall(rf"\b\w+\b", html_line) if w.lower() == wrong_lower), wrong_lower.title()
            )
            black = data["black"] or ""
            groq = data["groq"] or ""
            meaning = data["meaning"]

            span = (
                f"<span class='grammar-wrong' "
                f"data-wrong='{wrong_original}' "
                f"data-black='{black}' "
                f"data-groq='{groq}' "
                f"data-meaning='{meaning}'>{wrong_original}</span>"
            )
            html_line = re.sub(
                rf"\b{re.escape(wrong_original)}\b",
                span,
                html_line,
                count=1,
                flags=re.IGNORECASE
            )

        final_html.append(f"<p>{html_line}</p>")

    return "\n".join(final_html)

# -------------------------------------------------------
# 10) Routes
# -------------------------------------------------------
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        text_input = request.form.get("text", "").strip()
        file = request.files.get("file")

        if file and file.filename.endswith(".docx"):
            doc = Document(file)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            text = text_input

        if not text.strip():
            return render_template("index.html", error="No text found!")

        output = process_text_line_by_line(text)
        return render_template("result.html", highlighted_html=output)

    return render_template("index.html")

@app.route("/download_corrected", methods=["POST"])
def download_corrected():
    final_text = request.form.get("final_text", "")
    replacements = json.loads(request.form.get("replacements", "[]"))

    doc = Document()
    for line in final_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    # Apply replacements
    for para in doc.paragraphs:
        for rep in replacements:
            wrong = rep.get("old", "")
            correct = rep.get("new", "")
            if wrong and correct:
                para.text = re.sub(
                    rf"\b{re.escape(wrong)}\b",
                    correct,
                    para.text,
                    flags=re.IGNORECASE
                )

    output_path = "static/Corrected_Final_Output.docx"
    doc.save(output_path)

    return send_file(output_path, as_attachment=True, download_name="Corrected_Document.docx")

# -------------------------------------------------------
# 11) Run app
# -------------------------------------------------------
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
